import time
import cProfile
import pstats
import math
from io import StringIO
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt


# Функція для обчислення факторіалу числа
def calculate_factorial(n):
    time.sleep(1)
    return math.factorial(n)


# Функція для обчислення таблиці квадратів для ряду чисел
def calculate_squares(n):
    time.sleep(2)
    return [x ** 2 for x in range(n)]


# Функція для обчислення таблиці кубів для ряду чисел
def calculate_cubes(n):
    time.sleep(3)
    return [x ** 3 for x in range(n)]


# Функція для затримки на заданий проміжок часу
def delay(seconds):
    time.sleep(seconds)


# Основна функція, що викликає всі інші функції
def main():
    n = 1000  # приклад значення
    seconds = 1  # приклад затримки

    # Виклик функцій
    calculate_factorial(n)
    calculate_squares(n)
    calculate_cubes(n)
    delay(seconds)


# Профілювання основної функції
def profile_main():
    profiler = cProfile.Profile()
    profiler.enable()
    main()
    profiler.disable()
    return profiler


# Оптимізовані функції (з удвічі коротшою затримкою)
def optimized_calculate_factorial(n):
    time.sleep(0.5)
    return math.factorial(n)


def optimized_calculate_squares(n):
    time.sleep(1)
    return [x ** 2 for x in range(n)]


def optimized_calculate_cubes(n):
    time.sleep(1.5)
    return [x ** 3 for x in range(n)]


# Оновлена основна функція з оптимізованими функціями
def main_optimized():
    n = 1000  # приклад значення
    seconds = 0.5  # приклад затримки

    # Виклик функцій
    optimized_calculate_factorial(n)
    optimized_calculate_squares(n)
    optimized_calculate_cubes(n)
    delay(seconds)


# Профілювання оновленої основної функції
def profile_main_optimized():
    profiler = cProfile.Profile()
    profiler.enable()
    main_optimized()
    profiler.disable()
    return profiler


# Видалення повних шляхів до функцій в результатах профілювання
def extract_relevant_stats(stats):
    relevant_stats = {}
    function_names = ["calculate_factorial", "calculate_squares", "calculate_cubes", "delay"]

    for func_name in function_names:
        func_stats = next((stat for stat in stats.stats if func_name in stat[2]), None)
        if func_stats:
            relevant_stats[func_name] = stats.stats[func_stats][3]  # cumulative time
        else:
            relevant_stats[func_name] = 0

    return relevant_stats


# Створення діаграми часу виконання
def create_combined_plot(times_before, times_after, output_file):
    fig, ax = plt.subplots()
    function_names = ["calculate_factorial", "calculate_squares", "calculate_cubes", "delay"]

    bar_width = 0.35
    index = range(len(function_names))

    bar1 = ax.bar(index, [times_before[func] for func in function_names], bar_width, label='Before Optimization')
    bar2 = ax.bar([i + bar_width for i in index], [times_after[func] for func in function_names], bar_width,
                  label='After Optimization')

    ax.set_xlabel('Functions')
    ax.set_ylabel('Cumulative Time (seconds)')
    ax.set_title('Function Performance Before and After Optimization')
    ax.set_xticks([i + bar_width / 2 for i in index])
    ax.set_xticklabels(function_names)
    ax.legend()

    plt.tight_layout()
    plt.savefig(output_file)
    plt.close()


# Запис результатів у файл звіту (MS Word)
def create_report(plot_file, output_doc, times_before, times_after):
    doc = Document()
    doc.add_heading('Profile Analysis Report', 0)

    doc.add_heading('Performance Analysis', level=1)
    doc.add_picture(plot_file, width=Inches(6))

    doc.add_heading('Detailed Profile Before Optimization', level=2)
    for func, time in times_before.items():
        doc.add_paragraph(f"{func}: {time:.6f} seconds")

    doc.add_heading('Detailed Profile After Optimization', level=2)
    for func, time in times_after.items():
        doc.add_paragraph(f"{func}: {time:.6f} seconds")

    doc.save(output_doc)


# Основна функція для виконання всіх завдань
def run_all():
    # Профілювання не оптимізованого коду
    profiler_before = profile_main()
    stats_before = pstats.Stats(profiler_before).sort_stats('cumulative')
    times_before = extract_relevant_stats(stats_before)

    # Запис результатів профайлера у текстовий файл
    with open('profile_results.txt', 'w') as f:
        f.write("Profile results before optimization:\n")
        stats_before_output = StringIO()
        stats_before.stream = stats_before_output
        stats_before.print_stats()
        f.write(stats_before_output.getvalue())

    # Профілювання оптимізованого коду
    profiler_after = profile_main_optimized()
    stats_after = pstats.Stats(profiler_after).sort_stats('cumulative')
    times_after = extract_relevant_stats(stats_after)

    # Запис результатів профайлера у текстовий файл
    with open('profile_results.txt', 'a') as f:
        f.write("\nProfile results after optimization:\n")
        stats_after_output = StringIO()
        stats_after.stream = stats_after_output
        stats_after.print_stats()
        f.write(stats_after_output.getvalue())

    # Створення діаграми
    create_combined_plot(times_before, times_after, 'combined_plot.png')

    # Створення звіту
    create_report('combined_plot.png', 'Profile_Report.docx', times_before, times_after)


if __name__ == "__main__":
    run_all()
